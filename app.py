
from flask import Flask, request, jsonify, send_file
from docx import Document
from fpdf import FPDF
import os

app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

@app.route('/upload', methods=['POST'])
def upload_file():
    file = request.files.get('file')
    if file and file.filename.endswith('.docx'):
        filepath = os.path.join(UPLOAD_FOLDER, file.filename)
        file.save(filepath)

        # Extract metadata
        doc = Document(filepath)
        metadata = {
            "word_count": len(doc.paragraphs),
            "filename": file.filename
        }

        # Convert to PDF
        pdf_filepath = filepath.replace('.docx', '.pdf')
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font('Arial', size=12)
        for paragraph in doc.paragraphs:
            pdf.cell(200, 10, txt=paragraph.text, ln=True)
        pdf.output(pdf_filepath)

        return jsonify({
            "metadata": metadata,
            "pdf_url": f"/download/{os.path.basename(pdf_filepath)}"
        })

    return jsonify({"error": "Invalid file format"}), 400

@app.route('/download/<filename>', methods=['GET'])
def download_file(filename):
    return send_file(os.path.join(UPLOAD_FOLDER, filename))

if __name__ == '__main__':
    app.run(debug=True)
